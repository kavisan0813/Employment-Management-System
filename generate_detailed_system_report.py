import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

def create_report():
    doc = Document()

    # Setup margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    doc.add_paragraph()
    title = doc.add_paragraph("Employment Management System")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].bold = True

    subtitle = doc.add_paragraph("Comprehensive Technical & Architectural Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    date = doc.add_paragraph(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph("This document serves as a complete technical analysis of the Employment Management System. As requested for review by Team Leads, Project Managers, and Evaluation Panels, this report details every functional module, its underlying architecture, data flows, and future enhancements based exclusively on the actual source code.")
    
    # Features List to Iterate
    features = [
        {
            "name": "1. Authentication & Role-Based Access Control (RBAC)",
            "purpose": "To secure the application and ensure users only access data and features relevant to their roles.",
            "how_it_works": "The system utilizes React Context (AuthContext) and sessionStorage for state persistence. When a user logs in, their profile and role (Super Admin, HR Manager, Finance, Manager, Employee) are loaded into memory. Protected routes in 'routes.tsx' wrap page components with <Protected> and <RoleGuard> HOCs.",
            "user_flow": "1. User navigates to /login\n2. Enters credentials\n3. Clicks login\n4. Redirected to role-specific dashboard (e.g., HR sees HRDashboard, Employee sees Employee Dashboard).",
            "backend_flow": "Currently, login is verified against hardcoded demo accounts in AuthContext. No external server requests are made. In production, this would be an HTTP POST to /api/auth/login returning a JWT.",
            "db_operations": "Reads simulated user table to match email/password. Future state: SELECT query on 'users' table.",
            "apis": "AuthContext.login(), AuthContext.logout(), AuthContext.hasAccess()",
            "validations": "Checks if email and password fields are filled. Verifies credentials against local mock array.",
            "error_handling": "Displays 'Invalid credentials' toast notification on failure. Redirects unauthorized users accessing direct URLs back to /login."
        },
        {
            "name": "2. Dynamic Dashboard",
            "purpose": "Provide a high-level, at-a-glance view of organizational health, pending tasks, and key HR metrics customized to the user's role.",
            "how_it_works": "The Dashboard conditionally renders different components (AdminDashboard, HRDashboard, FinanceDashboard, ManagerDashboard) based on user.role. It aggregates data from the AppContext and renders Recharts visualizations.",
            "user_flow": "1. User logs in.\n2. Lands on /dashboard.\n3. Views aggregated KPI cards (Total Employees, Payroll, Pending Leaves).\n4. Clicks 'Quick Actions' to navigate to specific modules.",
            "backend_flow": "Currently aggregates mockData.ts arrays (employees, leaves, payroll) client-side. In production, an API endpoint (e.g., /api/dashboard/summary) would return pre-aggregated JSON payloads.",
            "db_operations": "Future state: Complex JOINs and GROUP BY queries aggregating counts from employees, leave_requests, and payroll tables.",
            "apis": "Internal Context getters: getEmployees(), getPendingLeaves()",
            "validations": "Ensures numeric data exists before rendering charts to prevent NaN errors.",
            "error_handling": "Recharts gracefully handles empty arrays by showing blank grids. Fallback UI provided if data is missing."
        },
        {
            "name": "3. Employee Management (CRUD)",
            "purpose": "Allow HR and Admins to manage the employee lifecycle, from adding new hires to updating roles and departments.",
            "how_it_works": "Accessed via the /employees route. Displays a searchable, filterable data table. Actions (Add, Edit, Delete) trigger Modals that update the global AppContext state.",
            "user_flow": "1. Admin navigates to Employees.\n2. Clicks 'Add Employee'.\n3. Fills out the modal form (Name, Role, Dept, Salary).\n4. Submits form.\n5. Table updates with new employee.",
            "backend_flow": "State is updated in AppContext arrays via React dispatch. In production: POST /api/employees to create, PUT /api/employees/:id to update, DELETE /api/employees/:id to remove.",
            "db_operations": "Future state: INSERT INTO employees (...); UPDATE employees SET (...) WHERE id=?; DELETE FROM employees WHERE id=?",
            "apis": "No external APIs. Relies on internal AppContext dispatch actions.",
            "validations": "Required fields (Name, Email, Role, Department). Email format validation.",
            "error_handling": "Displays toast errors if required fields are missing. Shows deletion confirmation dialogues to prevent accidental data loss."
        },
        {
            "name": "4. Recruitment Pipeline (ATS)",
            "purpose": "Track and manage job applicants through various hiring stages.",
            "how_it_works": "A Kanban-style board located at /recruitment. Uses HTML5 drag-and-drop mechanics (managed via a RecruitmentContext) to move candidates across columns: Applied, Screening, Round 1, Round 2, Offer, Hired.",
            "user_flow": "1. HR clicks Recruitment.\n2. Views candidates in stages.\n3. Drags a candidate from 'Applied' to 'Screening'.\n4. Drops the card; stage updates instantly.",
            "backend_flow": "Updates the 'stage' property of the candidate object in state. In production: PATCH /api/candidates/:id/stage.",
            "db_operations": "Future state: UPDATE candidates SET stage='Screening' WHERE id=?",
            "apis": "Internal state updates only. Uses contextual mock data models (Candidate, JobPosting).",
            "validations": "Prevents moving candidates to non-existent stages.",
            "error_handling": "Reverts drag-and-drop action if the drop zone is invalid."
        },
        {
            "name": "5. Leave & Request Management",
            "purpose": "Enable employees to request time off and allow managers/HR to approve or reject them.",
            "how_it_works": "Employees use a form to select dates and leave type. Managers view a consolidated list of pending requests for their direct reports. Statuses: Pending, Approved, Rejected.",
            "user_flow": "1. Employee navigates to Leave.\n2. Clicks 'Request Leave', picks dates & type.\n3. Submits. \n4. Manager navigates to Leave, sees request, clicks 'Approve'.",
            "backend_flow": "Appends to leaveRequests array in state. Manager action updates status. In production: POST /api/leaves to request, PATCH /api/leaves/:id to approve.",
            "db_operations": "Future state: INSERT INTO leave_requests; UPDATE leave_requests SET status='Approved' WHERE id=?",
            "apis": "Context methods for adding and updating leave requests.",
            "validations": "Start date must be before End date. Cannot request leave in the past. Requires a reason.",
            "error_handling": "Form validation prevents submission of invalid date ranges. Toast notification confirms successful submission."
        },
        {
            "name": "6. Attendance & Time Tracking",
            "purpose": "Log employee work hours and track punctuality.",
            "how_it_works": "Employees have a 'Clock In' / 'Clock Out' toggle. The system records timestamps. Admins can view historical logs.",
            "user_flow": "1. Employee clicks 'Clock In' button on dashboard.\n2. Timer starts.\n3. At end of day, clicks 'Clock Out'.",
            "backend_flow": "Generates a timestamp object in local state. In production: POST /api/attendance/clock-in with current UTC server time.",
            "db_operations": "Future state: INSERT INTO attendance (employee_id, check_in) VALUES (...); UPDATE attendance SET check_out=... WHERE ...",
            "apis": "None currently. Internal state array.",
            "validations": "Cannot clock out if not clocked in. Cannot clock in twice in the same shift.",
            "error_handling": "UI disables buttons based on current state (clocked in vs out)."
        },
        {
            "name": "7. Payroll Processing",
            "purpose": "Generate salary slips, calculate deductions, and track payments.",
            "how_it_works": "Restricted to Finance and Super Admin. Iterates over active employees, calculates gross vs net pay based on configured tax brackets/deductions, and generates digital payslips.",
            "user_flow": "1. Finance Admin navigates to Payroll.\n2. Selects Month/Year.\n3. Clicks 'Run Payroll'.\n4. System calculates amounts.\n5. Admin views individual Payslips in a modal.",
            "backend_flow": "A complex client-side mapping function computes values based on mock salary configurations. In production: POST /api/payroll/run which triggers a background worker (e.g., Celery/RabbitMQ) due to heavy calculation.",
            "db_operations": "Future state: Batch INSERT into payroll_records table linking to employee_id.",
            "apis": "None currently.",
            "validations": "Prevents running payroll twice for the same month/year.",
            "error_handling": "Alerts user if payroll for selected period already exists."
        }
    ]

    for f in features:
        doc.add_heading(f['name'], level=2)
        p = doc.add_paragraph()
        p.add_run("Purpose: ").bold = True
        p.add_run(f['purpose'] + "\n")
        
        p.add_run("How it works: ").bold = True
        p.add_run(f['how_it_works'] + "\n")
        
        p.add_run("User Flow: ").bold = True
        p.add_run("\n" + f['user_flow'] + "\n")
        
        p.add_run("Backend Flow: ").bold = True
        p.add_run(f['backend_flow'] + "\n")
        
        p.add_run("Database Operations: ").bold = True
        p.add_run(f['db_operations'] + "\n")
        
        p.add_run("APIs Involved: ").bold = True
        p.add_run(f['apis'] + "\n")
        
        p.add_run("Validations: ").bold = True
        p.add_run(f['validations'] + "\n")
        
        p.add_run("Error Handling: ").bold = True
        p.add_run(f['error_handling'] + "\n")
        
        doc.add_paragraph("_"*50)

    doc.add_page_break()

    # Database Schema representation
    doc.add_heading('Mock Database Schema (mockData.ts)', level=2)
    schema_data = [
        ["Table / Entity", "Key Fields", "Relationships"],
        ["Employees", "id, name, role, dept, salary", "Belongs to Dept"],
        ["Departments", "id, name, headId, budget", "Has many Employees"],
        ["Candidates", "id, name, stage, rating", "None"],
        ["LeaveRequests", "id, empId, dates, status", "Belongs to Employee"],
        ["Payroll", "id, empId, month, netPay", "Belongs to Employee"]
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i in range(3):
        hdr_cells[i].text = schema_data[0][i]
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    for row in schema_data[1:]:
        row_cells = table.add_row().cells
        for i in range(3):
            row_cells[i].text = row[i]

    doc.add_paragraph()

    # Enhancements
    doc.add_heading('Future Enhancements & Recommendations', level=2)
    doc.add_paragraph("Based on the code review, the following backend and infrastructural enhancements are required before a production release:")
    doc.add_paragraph("1. Database Implementation: Migrate from mockData.ts to a relational database (PostgreSQL) using an ORM like Prisma or Sequelize.", style='List Bullet')
    doc.add_paragraph("2. API Layer: Develop a Node.js/Express or Python/Django backend to handle CRUD operations and remove state persistence from the browser.", style='List Bullet')
    doc.add_paragraph("3. Authentication Security: Replace sessionStorage mock login with JWT-based authentication, bcrypt password hashing, and HttpOnly cookies.", style='List Bullet')
    doc.add_paragraph("4. Error Boundaries: Implement React Error Boundaries to catch unhandled JavaScript exceptions in components.", style='List Bullet')
    doc.add_paragraph("5. Testing: Introduce unit tests (Jest) for context reducers and E2E testing (Cypress/Playwright) for critical paths like Payroll and Recruitment.", style='List Bullet')
    
    output_path = 'NexusHR_Detailed_Technical_Report.docx'
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == '__main__':
    create_report()
