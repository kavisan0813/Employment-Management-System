from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style Definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── Helper Functions ──
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_flowchart(doc, title, steps):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"── {title} ──")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, step in enumerate(steps):
        cell = table.rows[0].cells[i]
        cell.text = step
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8)
                r.bold = True
    doc.add_paragraph()

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_para(doc, text="", bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

# ══════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NexusHR\nEmployment Management System")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Technical Analysis Report")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}\n"
                    f"Tech Stack: React 19 | TypeScript | Vite | Tailwind CSS\n"
                    f"Deployment: Netlify")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ── TABLE OF CONTENTS placeholder ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1.  Project Overview",
    "2.  Problem Statement",
    "3.  Objectives",
    "4.  Technology Stack",
    "5.  System Architecture",
    "6.  Folder Structure Explanation",
    "7.  Database Schema & Relationships",
    "8.  Authentication & Authorization Flow",
    "9.  Complete Feature List",
    "10. User Roles & Permissions",
    "11. API Documentation",
    "12. Frontend Functionality",
    "13. Backend Functionality",
    "14. Business Logic Workflows",
    "15. Dashboard & Reports Analysis",
    "16. Validations & Error Handling",
    "17. Security Features",
    "18. Performance Optimizations",
    "19. Third-Party Integrations",
    "20. Deployment Architecture",
    "21. Strengths",
    "22. Limitations",
    "23. Future Enhancements",
    "24. End-to-End User Flow",
    "25. Executive Summary"
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════
# SECTION 1: PROJECT OVERVIEW
# ══════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)
add_para(doc, (
    "NexusHR is a comprehensive, web-based Employment Management System designed to streamline and automate "
    "core HR operations for modern organizations. Built as a single-page application (SPA) using React 19 with "
    "TypeScript, NexusHR provides role-based access for five distinct user types: Super Admin, HR Manager, "
    "Finance, Manager, and Employee. The system covers the complete employee lifecycle from recruitment and "
    "onboarding through attendance tracking, leave management, expense management, payroll processing, "
    "performance reviews, training administration, shift scheduling, and advanced reporting."
))
add_para(doc, (
    "This prototype implementation uses a mock data layer (mockData.ts) containing 10 employees, 8 departments, "
    "6 leave requests, 5 expense claims, 7 recruitment candidates, and sample attendance/schedule records. "
    "All authentication is simulated via sessionStorage with hardcoded demo accounts per role. The project "
    "is deployed on Netlify with the build output in the dist/ directory."
))

# ══════════════════════════════════════════════════
# SECTION 2: PROBLEM STATEMENT
# ══════════════════════════════════════════════════
doc.add_heading('2. Problem Statement', level=1)
add_para(doc, (
    "Organizations managing employee data through fragmented systems — spreadsheets, email chains, and "
    "disconnected HR tools — face significant operational inefficiencies. These include:"
))
problems = [
    "No centralized repository for employee records, leading to data duplication and version conflicts.",
    "Manual leave request and approval workflows causing delays and tracking difficulties.",
    "Payroll computation errors due to disconnected attendance and leave data.",
    "Lack of real-time visibility into organizational KPIs for leadership decision-making.",
    "Inefficient recruitment pipeline tracking without standardized stage management.",
    "Difficulty coordinating training programs and tracking employee skill development.",
    "No unified expense reporting and reimbursement workflow.",
    "Role-based access control challenges when multiple stakeholders need different data views."
]
for p_text in problems:
    add_bullet(doc, p_text)

# ══════════════════════════════════════════════════
# SECTION 3: OBJECTIVES
# ══════════════════════════════════════════════════
doc.add_heading('3. Objectives', level=1)
objectives = [
    ("Centralized Data Hub: ", "Provide a single source of truth for all employee, department, and organizational data."),
    ("Role-Based Access: ", "Implement granular permissions ensuring each user sees only relevant data and actions."),
    ("Process Automation: ", "Digitize leave requests, expense claims, payroll runs, and recruitment pipelines through structured workflows."),
    ("Real-Time Dashboarding: ", "Deliver role-specific dashboards with key metrics, charts (Recharts), and actionable KPIs."),
    ("Employee Self-Service: ", "Empower employees to view payslips, submit requests, update profiles, and track attendance without HR intervention."),
    ("Comprehensive Reporting: ", "Provide multi-tab reporting with filtering, export, and trend analysis across all HR domains."),
    ("Modern UX: ", "Delight users with smooth animations (Framer Motion), dark/light theme, drag-and-drop Kanban boards, and a responsive layout."),
]
for prefix, text in objectives:
    add_bullet(doc, text, bold_prefix=prefix)

# ══════════════════════════════════════════════════
# SECTION 4: TECHNOLOGY STACK
# ══════════════════════════════════════════════════
doc.add_heading('4. Technology Stack', level=1)
add_table(doc,
    ["Category", "Technology", "Version", "Purpose"],
    [
        ["Framework", "React", "19.x", "UI component library & SPA architecture"],
        ["Language", "TypeScript", "~5.x", "Type-safe development & enhanced IDE support"],
        ["Build Tool", "Vite", "~6.x", "Fast dev server, HMR, optimized production builds"],
        ["Styling", "Tailwind CSS", "~4.x", "Utility-first CSS framework for rapid UI development"],
        ["Animations", "Framer Motion", "~12.x", "Declarative animations & gesture handling"],
        ["Charts", "Recharts", "~2.x", "Composable chart library for dashboards & reports"],
        ["UI Components", "Radix UI Primitives", "~1.x", "Accessible, unstyled UI primitives"],
        ["Icons", "Lucide React", "~0.x", "Consistent iconography throughout the app"],
        ["Notifications", "Sonner", "~2.x", "Toast notification system"],
        ["Routing", "React Router", "~7.x", "Client-side routing with lazy loading"],
        ["Deployment", "Netlify", "—", "Static site hosting with SPA redirect support"],
        ["Package Manager", "npm", "—", "Dependency management"],
        ["Date Handling", "date-fns", "~4.x", "Date formatting & manipulation"],
        ["Utilities", "clsx / tailwind-merge", "—", "Conditional class name merging"],
    ],
    col_widths=[3.5, 3.5, 2, 7]
)

# ══════════════════════════════════════════════════
# SECTION 5: SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════
doc.add_heading('5. System Architecture', level=1)
add_para(doc, (
    "NexusHR follows a component-based, layered architecture within the React SPA paradigm. The system "
    "is structured into four logical layers:"
))

add_para(doc, "Presentation Layer (Pages & Components)", bold=True)
add_bullet(doc, "60+ page components organized by domain (Payroll, Leave, Recruitment, etc.)")
add_bullet(doc, "Role-specific variants (ManagerLeaveApprovals, FinancePayroll, AdminDashboard)")
add_bullet(doc, "Shared UI library (button, card, badge, StatusBadge) in src/app/components/ui/")
add_bullet(doc, "Layout components (Sidebar, Topbar) with role-filtered navigation")

add_para(doc, "State Management Layer (Context Providers)", bold=True)
add_bullet(doc, "AuthContext: Authentication state, user role, hasAccess() permission checks")
add_bullet(doc, "AppContext: Application-wide state (theme, notifications, employee data)")
add_bullet(doc, "WorkflowContext: Recruitment pipeline stage management (drag-and-drop)")
add_bullet(doc, "React Context API used throughout — no external state library (Redux, Zustand)")

add_para(doc, "Data Layer (Mock Data)", bold=True)
add_bullet(doc, "mockData.ts: All entity data as typed arrays/objects")
add_bullet(doc, "No real backend, API, or database in this prototype")
add_bullet(doc, "Demonstrates the data shape for future backend integration")

add_para(doc, "Infrastructure Layer (Build & Deploy)", bold=True)
add_bullet(doc, "Vite dev server for development, optimized build for production")
add_bullet(doc, "Netlify deployment with netlify.toml configuration")
add_bullet(doc, "Code splitting via React.lazy() + Suspense for route-level chunking")

add_para(doc, "Architecture Diagram", bold=True)
add_flowchart(doc, "System Layers",
    ["User Browser", "React SPA", "Context API\nState Mgmt", "Mock Data\nLayer", "Netlify\nCDN"])

# ══════════════════════════════════════════════════
# SECTION 6: FOLDER STRUCTURE
# ══════════════════════════════════════════════════
doc.add_heading('6. Folder Structure Explanation', level=1)
add_para(doc, "The project follows a domain-organized structure within the src/ directory:")
structure = [
    ("src/app/", "Application core — routing, contexts, layouts, shared components"),
    ("src/app/components/ui/", "Reusable UI primitives (Button, Card, Badge, etc.)"),
    ("src/app/components/dashboards/", "Role-specific dashboard components (Admin, HR, Finance, Manager)"),
    ("src/app/components/", "Shared components (Sidebar, Topbar, StatusBadge, ToastNotification)"),
    ("src/app/context/", "React Context providers (AuthContext, AppContext, WorkflowContext)"),
    ("src/app/data/", "Mock data layer (mockData.ts with typed entity data)"),
    ("src/app/pages/", "Domain-specific page components (60+, role variants included)"),
    ("src/", "Entry points (main.tsx, App.tsx), global styles, Vite env config"),
    ("public/", "Static assets"),
    ("dist/", "Production build output (deployed to Netlify)"),
]
for path, desc in structure:
    add_bullet(doc, f"  {desc}", bold_prefix=path)

# ══════════════════════════════════════════════════
# SECTION 7: DATABASE SCHEMA & RELATIONSHIPS
# ══════════════════════════════════════════════════
doc.add_heading('7. Database Schema & Relationships', level=1)
add_para(doc, (
    "While NexusHR currently uses a mock data layer, the data models defined in mockData.ts represent the "
    "conceptual database schema. The following entities and relationships are defined:"
))

add_para(doc, "Entity Relationship Summary", bold=True)
add_table(doc,
    ["Entity", "Key Fields", "Relationships", "Sample Records"],
    [
        ["Employee", "id, name, email, departmentId, role, managerId, salary", "Belongs to Department; Reports to Manager", "10"],
        ["Department", "id, name, headId, budget", "Has many Employees", "8"],
        ["LeaveRequest", "id, employeeId, type, startDate, endDate, status", "Belongs to Employee", "6"],
        ["ExpenseClaim", "id, employeeId, category, amount, status", "Belongs to Employee", "5"],
        ["Attendance", "id, employeeId, date, status, checkIn, checkOut", "Belongs to Employee", "sample"],
        ["Candidate", "id, name, position, stage, appliedDate", "Independent (recruitment pipeline)", "7"],
        ["Training", "id, name, employeeIds[], startDate, status", "Many-to-Many with Employee", "sample"],
        ["Shift", "id, employeeId, date, startTime, endTime", "Belongs to Employee", "sample"],
        ["Payroll", "id, employeeId, month, year, grossPay, deductions, netPay, status", "Belongs to Employee", "sample"],
    ],
    col_widths=[2.5, 5, 4.5, 2.5]
)

add_para(doc, "Entity Relationship Diagram (Textual)", bold=True)
add_flowchart(doc, "Relationships",
    ["Department\n1", "── has many ──▶", "Employee\n*", "── has many ──▶", "LeaveRequest\n*"])
add_flowchart(doc, "",
    ["Employee\n*", "── has many ──▶", "ExpenseClaim\n*"])
add_flowchart(doc, "",
    ["Employee\n*", "── has many ──▶", "Attendance\n*"])
add_flowchart(doc, "",
    ["Employee\n*", "── many-to-many ──▶", "Training\n*"])

# ══════════════════════════════════════════════════
# SECTION 8: AUTHENTICATION & AUTHORIZATION FLOW
# ══════════════════════════════════════════════════
doc.add_heading('8. Authentication & Authorization Flow', level=1)
add_para(doc, "Authentication (Login Flow)", bold=True)
add_para(doc, (
    "Authentication is simulated in the browser using sessionStorage. There is no server-side auth. "
    "The login process works as follows:"
))
add_flowchart(doc, "Login Flow",
    ["User enters\ncredentials", "AuthContext\nvalidates", "sessionStorage\nstores user", "Redirect to\nhome route"])
add_para(doc, (
    "1. User submits email + password on the Login page\n"
    "2. AuthContext.login() checks against hardcoded demo accounts\n"
    "3. On success: user object and token are stored in sessionStorage\n"
    "4. AuthContext state updates, triggering route re-render\n"
    "5. User is redirected to their role-specific home route\n"
    "6. On failure: error message displayed via Sonner toast"
))

add_para(doc, "Authorization (Access Control)", bold=True)
add_para(doc, (
    "NexusHR implements a three-layer authorization model:"
))
add_para(doc)
add_para(doc, "Layer 1 — Route Guards (AuthGuard + RoleGuard)")
add_bullet(doc, "AuthGuard: Wraps protected routes; redirects to /login if unauthenticated")
add_bullet(doc, "RoleGuard: Wraps role-specific routes; restricts access based on user.role")
add_para(doc, "Layer 2 — hasAccess() Utility")
add_bullet(doc, "Defined in AuthContext, used throughout components to conditionally render UI elements")
add_bullet(doc, "Accepts an array of allowed roles, returns boolean")
add_para(doc, "Layer 3 — Role Wrappers & Conditional Rendering")
add_bullet(doc, "Role-specific page variants (e.g., AdminDashboard vs HRDashboard vs FinanceDashboard)")
add_bullet(doc, "Conditional rendering of sidebar items, buttons, and data columns based on role")

add_para(doc, "Role-Based Access Matrix", bold=True)
add_table(doc,
    ["Feature / Page", "Super Admin", "HR Manager", "Finance", "Manager", "Employee"],
    [
        ["Dashboard", "Admin View", "HR View", "Finance View", "Manager View", "Employee View"],
        ["Employee Management", "Full CRUD", "Full CRUD", "View Only", "Team View", "Self Only"],
        ["Payroll", "Full Access", "View Only", "Full Access", "Team View", "Own Payslip"],
        ["Leave Management", "All Requests", "All Requests", "View Only", "Team Approvals", "Own Requests"],
        ["Expense Management", "All Claims", "View Only", "Full Access", "Team Approvals", "Own Claims"],
        ["Recruitment", "Full Access", "Full Access", "No Access", "No Access", "No Access"],
        ["Performance Reviews", "All Reviews", "All Reviews", "No Access", "Team Reviews", "Own Review"],
        ["Training", "All Programs", "All Programs", "No Access", "Team View", "Own Enrollments"],
        ["Reports", "All Reports", "HR Reports", "Finance Reports", "Team Reports", "No Access"],
        ["Settings", "Full Access", "Limited", "Limited", "No Access", "Profile Only"],
    ],
    col_widths=[3.5, 2.2, 2.2, 2.2, 2.2, 2.2]
)

# ══════════════════════════════════════════════════
# SECTION 9: COMPLETE FEATURE LIST
# ══════════════════════════════════════════════════
doc.add_heading('9. Complete Feature List', level=1)
features = [
    ("Authentication & Authorization", "Login/Signup pages, session-based auth, role-based route guards, hasAccess() permission checks"),
    ("Dashboard (4 Role Variants)", "AdminDashboard, HRDashboard, FinanceDashboard, ManagerDashboard with KPIs & Recharts"),
    ("Employee Management", "Employee directory, profiles, department assignment, role management"),
    ("Leave Management", "Leave request/approval workflow, leave balance tracking, role-specific views (Employee, Manager, HR)"),
    ("Attendance Tracking", "Check-in/check-out records, attendance status, monthly attendance summary"),
    ("Expense Management", "Expense claim submission, approval workflow, category breakdowns, role-specific views"),
    ("Payroll Processing", "Payroll run workflow, payslip generation with modal view, month/year selection, earnings & deductions breakdown"),
    ("Recruitment Pipeline", "Kanban board with drag-and-drop (6 stages), candidate cards, interview scheduling, stage management"),
    ("Performance Reviews", "Review creation, rating system, goal setting, review history"),
    ("Training Management", "Training program creation, employee enrollment, completion tracking"),
    ("Shift & Schedule Management", "Shift assignment, weekly schedule view, shift swap requests"),
    ("Onboarding", "New hire onboarding checklist, task tracking, document upload simulation"),
    ("Employee Self-Service", "Profile viewing, payslip access, leave/expense request submission, attendance view"),
    ("Smart Search", "Global search across employees, departments, and records with filters"),
    ("Reports (Multi-Tab)", "8,533-line reporting interface with employee, payroll, leave, expense, attendance reports, plus export"),
    ("Dark/Light Theme", "Theme toggle with system preference detection, persisted to localStorage"),
    ("Toast Notifications", "Sonner-based toast system for success/error/info feedback"),
    ("Responsive Sidebar", "898-line role-filtered navigation with 28+ items, collapsible"),
    ("Topbar with Global Search", "952-line top navigation bar with global search, notifications, user menu, theme toggle"),
]
add_table(doc,
    ["Feature", "Description"],
    features,
    col_widths=[4.5, 11]
)

# ══════════════════════════════════════════════════
# SECTION 10: USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════
doc.add_heading('10. User Roles & Permissions', level=1)
add_para(doc, (
    "Five user roles are defined in AuthContext, each with a distinct permission scope "
    "and a role-specific home route:"
))
add_table(doc,
    ["Role", "Home Route", "Scope", "Permissions"],
    [
        ["Super Admin", "/dashboard", "Full system access", "All modules visible, all actions permitted"],
        ["HR Manager", "/dashboard", "HR operations", "Employee, leave, recruitment, training, performance — full CRUD except payroll edit"],
        ["Finance", "/dashboard", "Financial operations", "Payroll (full), expenses (full), reports (financial) — no recruitment/training access"],
        ["Manager", "/dashboard", "Team management", "Team member views, leave/expense approval for direct reports, performance reviews for team"],
        ["Employee", "/dashboard", "Self-service only", "Own profile, own payslip, own leave/expense requests, own attendance"],
    ],
    col_widths=[2.5, 2.5, 3.5, 7]
)

# ══════════════════════════════════════════════════
# SECTION 11: API DOCUMENTATION
# ══════════════════════════════════════════════════
doc.add_heading('11. API Documentation', level=1)
add_para(doc, (
    "As a frontend-only prototype, NexusHR does not expose or consume real REST APIs. "
    "All data operations are performed directly against the in-memory mockData.ts arrays. "
    "Below is the conceptual API contract that the frontend expects:"
))
add_table(doc,
    ["Entity", "Fetch (GET)", "Create (POST)", "Update (PUT)", "Delete (DELETE)"],
    [
        ["Employees", "/api/employees", "/api/employees", "/api/employees/:id", "/api/employees/:id"],
        ["Departments", "/api/departments", "/api/departments", "/api/departments/:id", "/api/departments/:id"],
        ["Leave Requests", "/api/leave-requests", "/api/leave-requests", "/api/leave-requests/:id/approve", "—"],
        ["Expenses", "/api/expenses", "/api/expenses", "/api/expenses/:id/approve", "—"],
        ["Attendance", "/api/attendance", "/api/attendance/check-in", "/api/attendance/:id", "—"],
        ["Payroll", "/api/payroll", "/api/payroll/run", "/api/payroll/:id", "—"],
        ["Candidates", "/api/candidates", "/api/candidates", "/api/candidates/:id/stage", "/api/candidates/:id"],
        ["Training", "/api/training", "/api/training", "/api/training/:id", "/api/training/:id"],
        ["Performance", "/api/reviews", "/api/reviews", "/api/reviews/:id", "/api/reviews/:id"],
        ["Auth", "/api/auth/login", "/api/auth/register", "—", "—"],
    ],
    col_widths=[3, 3.5, 3.5, 4, 3]
)
add_para(doc, (
    "Note: In the current prototype, these API calls are stubbed or replaced with direct mockData imports. "
    "A real implementation would replace mockData operations with fetch/axios calls to these endpoints."
))

# ══════════════════════════════════════════════════
# SECTION 12: FRONTEND FUNCTIONALITY
# ══════════════════════════════════════════════════
doc.add_heading('12. Frontend Functionality', level=1)
add_para(doc, (
    "The frontend is built entirely with React 19 and TypeScript. Key frontend capabilities include:"
))

add_para(doc, "Component Architecture", bold=True)
add_bullet(doc, "60+ page components, each implementing a distinct HR domain or workflow")
add_bullet(doc, "Reusable UI primitives (Button, Card, Badge, StatusBadge) with consistent styling")
add_bullet(doc, "4 role-specific dashboard components with contextual KPIs and Recharts visualizations")
add_bullet(doc, "Layout components (Sidebar: 898 lines, Topbar: 952 lines) with role-filtered content")

add_para(doc, "Routing & Code Splitting", bold=True)
add_bullet(doc, "React Router v7 for client-side routing (770-line routes.tsx)")
add_bullet(doc, "All page components lazy-loaded via React.lazy() + Suspense with fallback spinners")
add_bullet(doc, "Route guards: AuthGuard (authentication check) and RoleGuard (role-based restriction)")

add_para(doc, "State Management", bold=True)
add_bullet(doc, "AuthContext: Manages authentication state, user object, role config, isAuthenticated flag")
add_bullet(doc, "AppContext: Application-wide state including theme preference and notification data")
add_bullet(doc, "WorkflowContext: Manages recruitment pipeline stages and candidate transitions for drag-and-drop")

add_para(doc, "UI/UX Features", bold=True)
add_bullet(doc, "Dark/Light theme with system preference detection, persisted to localStorage")
add_bullet(doc, "Framer Motion animations for page transitions, modals, and interactive elements")
add_bullet(doc, "Recharts for dashboard charts (bar charts, pie charts, line charts)")
add_bullet(doc, "Sonner toast notifications for user feedback")
add_bullet(doc, "Drag-and-drop Kanban board for recruitment pipeline (WorkflowContext + HTML5 DnD)")

# ══════════════════════════════════════════════════
# SECTION 13: BACKEND FUNCTIONALITY
# ══════════════════════════════════════════════════
doc.add_heading('13. Backend Functionality', level=1)
add_para(doc, (
    "NexusHR currently has no backend implementation. All data operations are performed directly "
    "against the mock data layer. A production version would require:"
))
backends = [
    ("API Server: ", "Node.js/Express or Python/Django REST framework to serve CRUD endpoints"),
    ("Database: ", "PostgreSQL or MongoDB for persistent storage of all entity data"),
    ("Authentication Service: ", "JWT-based auth with refresh tokens, password hashing (bcrypt), and session management"),
    ("File Storage: ", "Cloud storage (AWS S3, Cloudinary) for document uploads (receipts, resumes, payslips)"),
    ("Email Service: ", "Transactional emails for leave approvals, payroll notifications, password resets"),
    ("Background Jobs: ", "Scheduled payroll runs, attendance report generation, automated reminders"),
]
for prefix, text in backends:
    add_bullet(doc, text, bold_prefix=prefix)

add_para(doc, "Database Schema (Target)", bold=True)
add_table(doc,
    ["Table", "Key Columns", "Indexes", "Foreign Keys"],
    [
        ["employees", "id, name, email, role, department_id, manager_id, salary", "email (UNIQUE), department_id", "department_id → departments.id"],
        ["departments", "id, name, head_id, budget", "name (UNIQUE)", "head_id → employees.id"],
        ["leave_requests", "id, employee_id, type, start_date, end_date, status", "employee_id, status", "employee_id → employees.id"],
        ["expenses", "id, employee_id, category, amount, status, receipt_url", "employee_id, status", "employee_id → employees.id"],
        ["attendance", "id, employee_id, date, status, check_in, check_out", "employee_id, date", "employee_id → employees.id"],
        ["payroll", "id, employee_id, month, year, gross_pay, net_pay, status", "employee_id, month, year", "employee_id → employees.id"],
        ["candidates", "id, name, position, stage, email, phone", "email (UNIQUE)", "—"],
        ["training_programs", "id, name, start_date, end_date, status", "status", "—"],
        ["training_enrollments", "id, employee_id, program_id, status", "employee_id, program_id", "employee_id → employees.id, program_id → training_programs.id"],
        ["performance_reviews", "id, employee_id, reviewer_id, rating, period", "employee_id, reviewer_id", "employee_id → employees.id, reviewer_id → employees.id"],
    ],
    col_widths=[3, 5.5, 3.5, 5]
)

# ══════════════════════════════════════════════════
# SECTION 14: BUSINESS LOGIC WORKFLOWS
# ══════════════════════════════════════════════════
doc.add_heading('14. Business Logic Workflows', level=1)

add_para(doc, "Workflow 1: Leave Request & Approval", bold=True)
add_flowchart(doc, "Leave Approval Flow",
    ["Employee\nsubmits leave", "Manager\nreviews", "HR\nvalidates balance", "Approved /\nRejected", "Employee\nnotified"])
add_bullet(doc, "Employee selects leave type, dates, and submits reason")
add_bullet(doc, "Manager sees pending team requests in ManagerLeaveApprovals page")
add_bullet(doc, "Manager approves or rejects; status updates in real-time")
add_bullet(doc, "HR can override or manage all requests from LeaveManagement page")
add_bullet(doc, "Leave balance is updated upon approval (mock data only)")

add_para(doc, "Workflow 2: Payroll Processing", bold=True)
add_flowchart(doc, "Payroll Run Flow",
    ["Admin selects\nmonth/year", "System loads\nemployees", "Deductions\ncomputed", "Payslips\ngenerated", "Marked as\nPaid"])
add_bullet(doc, "Finance/Super Admin selects month and year")
add_bullet(doc, "System computes gross pay, deductions (tax, insurance, retirement), and net pay")
add_bullet(doc, "Payslips are generated with detailed breakdown")
add_bullet(doc, "Individual payslips viewable in modal with print-ready layout")
add_bullet(doc, "Status: Pending → Processing → Completed → Paid")

add_para(doc, "Workflow 3: Recruitment Pipeline", bold=True)
add_flowchart(doc, "Recruitment Pipeline",
    ["Sourced", "Applied", "Screened", "Interview", "Offer", "Hired"])
add_bullet(doc, "Candidates move through 6 stages via drag-and-drop on Kanban board")
add_bullet(doc, "Each candidate card shows name, position, applied date, skills")
add_bullet(doc, "Stage transitions update WorkflowContext state")
add_bullet(doc, "Interview scheduling dialog available from candidate cards")

add_para(doc, "Workflow 4: Expense Claim Processing", bold=True)
add_flowchart(doc, "Expense Approval Flow",
    ["Employee\nsubmits claim", "Manager\napproves", "Finance\nreviews", "Reimbursed /\nRejected", "Employee\nnotified"])
add_bullet(doc, "Employee submits claim with category, amount, and description")
add_bullet(doc, "Manager reviews and approves/rejects team claims")
add_bullet(doc, "Finance processes approved claims and marks as reimbursed")
add_bullet(doc, "Expense categories: Travel, Office Supplies, Meals, Accommodation, Transportation, Other")

# ══════════════════════════════════════════════════
# SECTION 15: DASHBOARD & REPORTS ANALYSIS
# ══════════════════════════════════════════════════
doc.add_heading('15. Dashboard & Reports Analysis', level=1)

add_para(doc, "Role-Specific Dashboards", bold=True)
add_table(doc,
    ["Role", "KPIs Displayed", "Charts", "Widgets"],
    [
        ["Super Admin", "Total employees, active leaves, pending expenses, open positions", "Bar chart (department headcount), Pie chart (leave types)", "Recent activity feed, quick actions"],
        ["HR Manager", "Employee count, leave utilization, training completion, new hires", "Training progress bar chart, department distribution", "Upcoming reviews, pendingonboarding"],
        ["Finance", "Payroll total, pending reimbursements, expense totals, budget utilization", "Expense by category pie, payroll trend line", "Month-over-month comparison"],
        ["Manager", "Team size, pending approvals, attendance rate, upcoming leaves", "Team attendance bar chart, leave distribution", "Team member quick view"],
    ],
    col_widths=[3, 4, 4, 4.5]
)

add_para(doc, "Reports Module (src/app/pages/Reports.tsx — 8,533 lines)", bold=True)
add_bullet(doc, "Multi-tab interface with separate report categories")
add_bullet(doc, "Employee Reports: Headcount, demographics, department distribution")
add_bullet(doc, "Payroll Reports: Salary summaries, deduction breakdowns, month-to-month trends")
add_bullet(doc, "Leave Reports: Utilization rates, type breakdown, department comparison")
add_bullet(doc, "Expense Reports: Category analysis, approval timelines, reimbursement status")
add_bullet(doc, "Attendance Reports: Present/absent rates, late arrivals, overtime tracking")
add_bullet(doc, "Export functionality for each report (CSV/PDF conceptual)")

# ══════════════════════════════════════════════════
# SECTION 16: VALIDATIONS & ERROR HANDLING
# ══════════════════════════════════════════════════
doc.add_heading('16. Validations & Error Handling', level=1)

add_para(doc, "Form Validations", bold=True)
add_bullet(doc, "Login/Signup: Email format validation, password length requirements, required field checks")
add_bullet(doc, "Leave Request: Date range validation (start < end), future date enforcement, required reason")
add_bullet(doc, "Expense Claim: Amount > 0, category selection required, description required")
add_bullet(doc, "Employee Creation: Required fields (name, email, role, department), email format")
add_bullet(doc, "Payroll: Month/year selection validation, duplicate run prevention (conceptual)")

add_para(doc, "Error Handling Patterns", bold=True)
add_bullet(doc, "Sonner toasts for user-facing success/error/info messages")
add_bullet(doc, "Login errors: Invalid credentials message displayed via toast")
add_bullet(doc, "Form submission errors: Validation feedback via inline messages and toasts")
add_bullet(doc, "Route guard redirects: Unauthorized users redirected to login")
add_bullet(doc, "Fallback UI: Suspense fallback spinners during lazy-loaded route transitions")

# ══════════════════════════════════════════════════
# SECTION 17: SECURITY FEATURES
# ══════════════════════════════════════════════════
doc.add_heading('17. Security Features', level=1)
add_para(doc, "Implemented:", bold=True)
add_bullet(doc, "Role-based access control at both route level (RoleGuard) and component level (hasAccess())")
add_bullet(doc, "Private routes protected by AuthGuard — unauthenticated users cannot access internal pages")
add_bullet(doc, "Role-specific home routes preventing post-login access to wrong dashboard")
add_bullet(doc, "Conditional UI rendering — unauthorized users never see restricted actions or data")
add_bullet(doc, "Session-based auth stored in sessionStorage (cleared on tab close)")

add_para(doc, "Conceptual (for production):", bold=True)
add_bullet(doc, "JWT-based authentication with access + refresh tokens")
add_bullet(doc, "Password hashing (bcrypt) and HTTPS enforcement")
add_bullet(doc, "Input sanitization and CSRF protection on API endpoints")
add_bullet(doc, "Rate limiting on login endpoints")
add_bullet(doc, "Audit logging for sensitive operations (payroll runs, data deletion)")

# ══════════════════════════════════════════════════
# SECTION 18: PERFORMANCE OPTIMIZATIONS
# ══════════════════════════════════════════════════
doc.add_heading('18. Performance Optimizations', level=1)
add_bullet(doc, "Code Splitting: All page components lazy-loaded via React.lazy() + Suspense, reducing initial bundle size")
add_bullet(doc, "Vite Build: Fast production builds with tree-shaking, minification, and chunk splitting")
add_bullet(doc, "TypeScript: Static type checking prevents runtime type errors")
add_bullet(doc, "Tailwind CSS: Utility-first approach with purging of unused styles in production")
add_bullet(doc, "Component Reusability: Shared UI primitives reduce code duplication")
add_bullet(doc, "Netlify CDN: Global edge caching for static assets")
add_bullet(doc, "Framer Motion: Optimized animations with hardware-accelerated CSS transforms")

# ══════════════════════════════════════════════════
# SECTION 19: THIRD-PARTY INTEGRATIONS
# ══════════════════════════════════════════════════
doc.add_heading('19. Third-Party Integrations', level=1)
add_table(doc,
    ["Library", "Version", "Purpose", "Usage Location"],
    [
        ["react", "^19.0.0", "UI framework", "All components"],
        ["react-dom", "^19.0.0", "DOM rendering", "main.tsx entry point"],
        ["typescript", "~5.7.0", "Type safety", "All .ts/.tsx files"],
        ["vite", "^6.0.0", "Build tool & dev server", "Build pipeline"],
        ["tailwindcss", "^4.0.0", "CSS framework", "All styling"],
        ["@tailwindcss/vite", "^4.0.0", "Tailwind Vite plugin", "vite.config.ts"],
        ["react-router", "^7.0.0", "Client-side routing", "routes.tsx, App.tsx"],
        ["framer-motion", "^12.0.0", "Animations", "Page transitions, modals, drag-and-drop"],
        ["recharts", "^2.0.0", "Charts & data visualization", "Dashboards, Reports"],
        ["lucide-react", "^0.0.0", "Icon set", "Sidebar, Topbar, buttons, indicators"],
        ["sonner", "^2.0.0", "Toast notifications", "Auth, form submissions, actions"],
        ["date-fns", "^4.0.0", "Date utilities", "Date formatting throughout"],
        ["clsx", "^2.0.0", "Conditional CSS classes", "Component styling"],
        ["tailwind-merge", "^3.0.0", "Tailwind class merging", "utils.ts"],
        ["@radix-ui/react-dialog", "^1.0.0", "Accessible dialogs/modals", "Modals (payslip, etc.)"],
        ["@radix-ui/react-dropdown-menu", "^1.0.0", "Accessible dropdowns", "User menu, filters"],
        ["@radix-ui/react-select", "^1.0.0", "Accessible selects", "Dropdown selects in forms"],
        ["@radix-ui/react-tabs", "^1.0.0", "Accessible tabs", "Reports tab interface"],
    ],
    col_widths=[4, 2, 4.5, 5.5]
)

# ══════════════════════════════════════════════════
# SECTION 20: DEPLOYMENT ARCHITECTURE
# ══════════════════════════════════════════════════
doc.add_heading('20. Deployment Architecture', level=1)
add_para(doc, "Current Deployment (Netlify)", bold=True)
add_bullet(doc, "Hosting: Netlify static site hosting with global CDN")
add_bullet(doc, "Build Command: npm run build (Vite production build)")
add_bullet(doc, "Publish Directory: dist/ (Vite output)")
add_bullet(doc, "SPA Redirect: netlify.toml configured to redirect all routes to index.html for client-side routing")
add_bullet(doc, "Environment Variables: Managed via Netlify dashboard")

add_para(doc, "Deployment Pipeline (Current)", bold=True)
add_flowchart(doc, "Deployment Flow",
    ["Git Push", "Netlify\nauto-detects", "npm install", "npm run build", "dist/\ndistributed\nvia CDN"])
add_para(doc)
add_para(doc, "Target Production Architecture", bold=True)
add_flowchart(doc, "Production Architecture",
    ["User\nBrowser", "Netlify CDN\n(Static Assets)", "API Gateway\n(Cloudflare/AWS)", "Backend\nServer", "Database\n(PostgreSQL)"])

# ══════════════════════════════════════════════════
# SECTION 21: STRENGTHS
# ══════════════════════════════════════════════════
doc.add_heading('21. Strengths', level=1)
strengths = [
    "Comprehensive Domain Coverage: Covers the full employee lifecycle — recruitment, onboarding, attendance, leave, expenses, payroll, performance, training, and reporting — all in a single application.",
    "Role-Based Architecture: Five well-defined roles with granular permissions at both route and component levels, ensuring appropriate data access for each user type.",
    "Modern Tech Stack: React 19 + TypeScript + Vite provides excellent developer experience, type safety, and fast build times. Tailwind CSS enables rapid, consistent UI development.",
    "Rich UI/UX: Dark/light themes, smooth animations (Framer Motion), interactive charts (Recharts), drag-and-drop Kanban (Recruitment), and toast notifications create a polished user experience.",
    "Code Organization: Clear domain-based file structure with 60+ page components, shared UI primitives, and role-specific dashboard variants makes the codebase navigable.",
    "Lazy Loading & Code Splitting: All routes are lazy-loaded, reducing initial bundle size and improving time-to-interactive.",
    "Comprehensive Reporting: The 8,533-line Reports module demonstrates deep commitment to data-driven decision making with multi-tab analysis across all HR domains.",
    "Accessibility: Use of Radix UI primitives ensures accessible dialog, dropdown, select, and tab components out of the box.",
]
for s in strengths:
    add_bullet(doc, s)

# ══════════════════════════════════════════════════
# SECTION 22: LIMITATIONS
# ══════════════════════════════════════════════════
doc.add_heading('22. Limitations', level=1)
limitations = [
    "No Backend: All data is mock data stored in-memory (mockData.ts). No persistence, no real API, no database. All data resets on page refresh.",
    "No Real Authentication: Login is simulated via sessionStorage with hardcoded demo accounts. No password hashing, JWT tokens, or session management.",
    "No Data Persistence: Changes made during a session (leaves approved, expenses submitted, payroll runs) are lost on page reload.",
    "Limited Mock Data: Only 10 employees, 8 departments, 6 leave requests — insufficient for realistic load testing or demonstrating scale.",
    "No File Upload: While expense claims conceptually support receipt uploads, no actual file upload or storage mechanism exists.",
    "No Email Notifications: There is no email service. Users are not notified of approvals, payroll completion, or other events.",
    "No Audit Trail: Sensitive operations (payroll runs, data changes) are not logged for compliance purposes.",
    "No Testing: The project lacks unit tests, integration tests, or end-to-end tests, making refactoring risky.",
    "Large Component Files: Reports.tsx (8,533 lines) and other large files violate the single-responsibility principle and hinder maintainability.",
    "No Real API Contracts: The frontend directly imports mock data rather than using fetch/axios, meaning all data access patterns would need rewriting for a real backend.",
    "No Error Boundaries: The application lacks React error boundaries, meaning an uncaught error in one component can crash the entire app.",
]
for l in limitations:
    add_bullet(doc, l)

# ══════════════════════════════════════════════════
# SECTION 23: FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════
doc.add_heading('23. Future Enhancements', level=1)
enhancements = [
    ("Backend Integration: ", "Replace mockData.ts with a proper backend (Node.js/Express or Python/Django) with RESTful API endpoints and a PostgreSQL database."),
    ("Real Authentication: ", "Implement JWT-based authentication with login/register endpoints, password hashing (bcrypt), refresh tokens, and session management."),
    ("Data Persistence: ", "Integrate a production database to persist all employee, leave, expense, payroll, and recruitment data."),
    ("File Management: ", "Add file upload support (receipts, resumes, payslips) using cloud storage (AWS S3, Cloudinary) with preview capabilities."),
    ("Email & Notifications: ", "Integrate an email service (SendGrid, AWS SES) for transactional notifications — leave approvals, payroll alerts, password resets."),
    ("Audit Logging: ", "Implement detailed audit trails for all sensitive operations to support compliance requirements."),
    ("Testing Suite: ", "Add unit tests (Vitest/Jest), integration tests, and E2E tests (Playwright/Cypress) to ensure reliability."),
    ("CI/CD Pipeline: ", "Set up GitHub Actions for automated testing, linting, type-checking, and deployment to staging/production environments."),
    ("Performance Optimization: ", "Implement virtualization (react-window) for large lists, memoization (useMemo/useCallback) for expensive computations, and bundle analysis."),
    ("Multi-Language Support: ", "Add i18n support (react-intl, i18next) for internationalization."),
    ("Mobile App: ", "Develop companion mobile apps (React Native) for attendance check-in, leave requests, and notifications."),
    ("Advanced Analytics: ", "Integrate AI/ML for predictive analytics — attrition risk prediction, hiring demand forecasting, budget optimization."),
    ("Document Generation: ", "Auto-generate offer letters, employment contracts, and performance reports as PDFs."),
    ("Time Tracking: ", "Add project-based time tracking with timesheets, billable hours, and client invoicing integration."),
]
for prefix, text in enhancements:
    add_bullet(doc, text, bold_prefix=prefix)

# ══════════════════════════════════════════════════
# SECTION 24: END-TO-END USER FLOW
# ══════════════════════════════════════════════════
doc.add_heading('24. End-to-End User Flow', level=1)

add_para(doc, "Scenario: Employee submits a leave request and it gets approved", bold=True)
add_flowchart(doc, "E2E Leave Request Flow",
    ["Login", "Dashboard\n(Employee View)", "Leave\nManagement", "Submit\nRequest", "Logout"])
add_flowchart(doc, "",
    ["Manager\nLogin", "Dashboard\n(Manager View)", "Pending\nApprovals", "Approve\nRequest", "Status\nUpdated"])
add_para(doc)
steps_e2e = [
    "1. Employee navigates to the app URL (https://nexushr.netlify.app)",
    "2. Login page loads — employee enters demo credentials",
    "3. AuthContext validates credentials, stores user in sessionStorage, redirects to Employee Dashboard",
    "4. Employee Dashboard shows personal KPIs: attendance rate, pending leaves, next pay date",
    "5. Employee clicks 'Leave Management' in the sidebar",
    "6. LeaveManagement page loads — shows employee's leave history and remaining balance",
    "7. Employee clicks 'Apply Leave' — form appears with leave type, dates, reason fields",
    "8. Employee fills form, submits — validation runs, success toast appears",
    "9. Leave request is added to mockData with status 'Pending'",
    "10. Employee logs out (sessionStorage cleared, redirected to login)",
    "",
    "11. Manager logs in with different credentials — redirected to Manager Dashboard",
    "12. Manager sees 'Pending Approvals' KPI on dashboard showing 1 pending request",
    "13. Manager clicks Leave Management → sees team leave requests including employee's submission",
    "14. Manager clicks 'Approve' — confirmation dialog appears",
    "15. On confirmation, leave status updates to 'Approved', success toast appears",
    "16. Leave balance is updated in mock data",
    "",
    "17. Employee logs back in → Leave Management shows approved status",
    "18. Dashboard now shows updated leave balance",
]
for step in steps_e2e:
    p = doc.add_paragraph(step)
    p.paragraph_format.space_after = Pt(1)

# ══════════════════════════════════════════════════
# SECTION 25: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════
doc.add_heading('25. Executive Summary', level=1)
add_para(doc, (
    "NexusHR is a modern, feature-rich Employment Management System prototype that demonstrates "
    "a comprehensive approach to digitizing HR operations. Built with React 19, TypeScript, and "
    "a modern toolchain (Vite + Tailwind CSS), the application covers nine major HR domains across "
    "60+ page components with role-based access for five user types."
))
add_para(doc, (
    "The system's architecture is organized into clear layers: a presentation layer of domain-specific "
    "pages and shared UI components, a state management layer using React Context (AuthContext, AppContext, "
    "WorkflowContext), and a mock data layer that defines the complete data schema for all entities. "
    "Role-based access control is implemented at both the route level (AuthGuard, RoleGuard) and component "
    "level (hasAccess() utility), providing granular permission management."
))
add_para(doc, (
    "Key strengths include comprehensive domain coverage (recruitment through payroll and reporting), "
    "a modern and polished UI with dark/light theme support, smooth animations, interactive dashboards "
    "with Recharts, and a Kanban-style recruitment pipeline with drag-and-drop. The project demonstrates "
    "strong code organization principles with lazy-loaded routes, shared component libraries, and "
    "role-specific page variants."
))
add_para(doc, (
    "As a prototype, the primary limitations are the absence of a backend server, database persistence, "
    "real authentication, file uploads, email notifications, and automated testing. These limitations "
    "represent clear next steps for production readiness. The well-structured codebase, comprehensive "
    "data models, and clearly defined access patterns provide a solid foundation for backend integration "
    "and feature expansion."
))
add_para(doc, (
    "NexusHR successfully demonstrates the vision of an all-in-one HR management platform with a "
    "user-centric design, role-appropriate interfaces, and data-driven decision support. The project "
    "is well-positioned for evolution from a frontend prototype to a full-stack production application."
))

# ══════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════
output_path = "D:\\Employment Management System\\NexusHR_Technical_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
